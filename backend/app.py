import os
import io
import csv
import json
import traceback
import subprocess
from datetime import datetime, timedelta

# Flask & Extensions
from flask import Flask, jsonify, request, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import (
    create_access_token, jwt_required, get_jwt_identity, 
    get_jwt, JWTManager
)
from flask_migrate import Migrate
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# SQLAlchemy
from sqlalchemy import case, or_
from sqlalchemy.orm import joinedload
from sqlalchemy.sql import func

# Data & Documents
import pandas as pd
from docx import Document
from pypdf import PdfWriter
import barcode
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
)
from reportlab.graphics.barcode import code128

# ==============================================================================
# CONFIGURAÇÃO INICIAL
# ==============================================================================

app = Flask(__name__)

# Configuração JWT
app.config["JWT_SECRET_KEY"] = "senha_padrao_dev" 
jwt = JWTManager(app)

# Configuração Banco de Dados

DB_USER = os.getenv('DB_USER', 'root')
DB_PASS = os.getenv('DB_PASS', 'senha_padrao_dev') # Senha genérica
DB_HOST = os.getenv('DB_HOST', 'localhost')        # Localhost para quem baixar
DB_NAME = os.getenv('DB_NAME', 'estoque_db')

app.config['SQLALCHEMY_DATABASE_URI'] = f'mysql+pymysql://{DB_USER}:{DB_PASS}@{DB_HOST}/{DB_NAME}'app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
migrate = Migrate(app, db)

# ==============================================================================
# TABELAS DE ASSOCIAÇÃO
# ==============================================================================

produto_fornecedor = db.Table('produto_fornecedor',
    db.Column('FK_PRODUTO_Id_produto', db.Integer, db.ForeignKey('produto.Id_produto'), primary_key=True),
    db.Column('FK_FORNECEDOR_id_fornecedor', db.Integer, db.ForeignKey('fornecedor.id_fornecedor'), primary_key=True)
)

produto_natureza = db.Table('produto_natureza',
    db.Column('fk_PRODUTO_Id_produto', db.Integer, db.ForeignKey('produto.Id_produto'), primary_key=True),
    db.Column('fk_NATUREZA_id_natureza', db.Integer, db.ForeignKey('natureza.id_natureza'), primary_key=True)
)

# ==============================================================================
# MODELOS (ENTITIES)
# ==============================================================================

class Produto(db.Model):
    __tablename__ = 'produto'
    id_produto = db.Column('Id_produto', db.Integer, primary_key=True)
    nome = db.Column('Nome', db.String(100), nullable=False)
    codigo = db.Column('Codigo', db.String(20), unique=True, nullable=False)
    descricao = db.Column('Descricao', db.String(200))
    preco = db.Column('Preco', db.Numeric(10, 2), nullable=True, default=0.00)
    codigoB = db.Column('CodigoB', db.String(20))
    codigoC = db.Column('CodigoC', db.String(20))
    
    fornecedores = db.relationship('Fornecedor', secondary=produto_fornecedor, back_populates='produtos')
    naturezas = db.relationship('Natureza', secondary=produto_natureza, back_populates='produtos')

class Fornecedor(db.Model):
    __tablename__ = 'fornecedor'
    id_fornecedor = db.Column(db.Integer, primary_key=True)
    nome = db.Column('Nome', db.String(50), unique=True, nullable=False)
    produtos = db.relationship('Produto', secondary=produto_fornecedor, back_populates='fornecedores')

class Natureza(db.Model):
    __tablename__ = 'natureza'
    id_natureza = db.Column(db.Integer, primary_key=True)
    nome = db.Column('nome', db.String(100), unique=True, nullable=False)
    produtos = db.relationship('Produto', secondary=produto_natureza, back_populates='naturezas')

class MovimentacaoEstoque(db.Model):
    __tablename__ = 'mov_estoque'
    id_movimentacao = db.Column(db.Integer, primary_key=True)
    id_produto = db.Column(db.Integer, db.ForeignKey('produto.Id_produto'), nullable=False)
    id_usuario = db.Column(db.Integer, db.ForeignKey('usuario.id_usuario'), nullable=False)
    data_hora = db.Column(db.DateTime, nullable=False, default=datetime.now)
    quantidade = db.Column(db.Integer, nullable=False)
    tipo = db.Column(db.Enum("Entrada", "Saida"), nullable=False)
    motivo_saida = db.Column(db.String(200))
    
    produto = db.relationship('Produto')
    usuario = db.relationship('Usuario')

class Usuario(db.Model):
    __tablename__ = 'usuario'
    id_usuario = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(100), nullable=False)
    login = db.Column(db.String(100), unique=True, nullable=False)
    senha_hash = db.Column(db.String(255), nullable=False)
    permissao = db.Column(db.String(100), nullable=False)
    ativo = db.Column(db.Boolean, default=True, nullable=False)

    def set_password(self, senha):
        self.senha_hash = generate_password_hash(senha)

    def check_password(self, senha):
        return check_password_hash(self.senha_hash, senha)

class Servico(db.Model):
    __tablename__ = 'servico'
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(200), nullable=False, unique=True)
    descricao = db.Column(db.Text, nullable=True)
    documentos = db.relationship('DocumentosGerados', back_populates='servico', cascade="all, delete-orphan")

    def __repr__(self):
        return f'<Servico {self.nome}>'
    
class DocumentosGerados(db.Model):
    __tablename__ = 'documentos_gerados'
    id = db.Column(db.Integer, primary_key=True)
    servico_id = db.Column(db.Integer, db.ForeignKey('servico.id'), nullable=False)
    usuario_id = db.Column(db.Integer, db.ForeignKey('usuario.id_usuario'), nullable=False)
    data_criacao = db.Column(db.DateTime, nullable=False, default=datetime.now)
    versao = db.Column(db.Integer, nullable=False)
    dados_formulario = db.Column(db.JSON, nullable=False)
    caminho_pdf_final = db.Column(db.String(255), nullable=False)
    
    servico = db.relationship('Servico', back_populates='documentos')
    usuario = db.relationship('Usuario')

    def __repr__(self):
        return f'<Documento v{self.versao} para Serviço ID {self.servico_id}>'

# ==============================================================================
# FUNÇÕES AUXILIARES
# ==============================================================================

def calcular_saldo_produto(id_produto):
    saldo = db.session.query(
        db.func.sum(
            case(
                (MovimentacaoEstoque.tipo == 'Entrada', MovimentacaoEstoque.quantidade),
                (MovimentacaoEstoque.tipo == 'Saida', -MovimentacaoEstoque.quantidade)
            )
        )
    ).filter(MovimentacaoEstoque.id_produto == id_produto).scalar() or 0
    return saldo

# ==============================================================================
# ROTAS: PRODUTOS
# ==============================================================================

@app.route('/api/produtos', methods=['GET'])
@jwt_required()
def get_todos_produtos():
    try:
        termo_busca = request.args.get('search')
        query = Produto.query
        
        if termo_busca:
            query = query.filter(
                or_(
                    Produto.nome.ilike(f"%{termo_busca}%"),
                    Produto.codigo.ilike(f"%{termo_busca}%"),
                    Produto.codigoB.ilike(f"%{termo_busca}%"),
                    Produto.codigoC.ilike(f"%{termo_busca}%")
                )
            )
        produtos_db = query.all()
        
        if not produtos_db:
            return jsonify([]), 200

        product_ids = [p.id_produto for p in produtos_db]
        fornecedores_map = {f.id_fornecedor: f.nome for f in Fornecedor.query.all()}
        naturezas_map = {n.id_natureza: n.nome for n in Natureza.query.all()}
        
        prod_forn_assoc = db.session.query(produto_fornecedor).filter(produto_fornecedor.c.FK_PRODUTO_Id_produto.in_(product_ids)).all()
        prod_nat_assoc = db.session.query(produto_natureza).filter(produto_natureza.c.fk_PRODUTO_Id_produto.in_(product_ids)).all()

        produto_fornecedores = {}
        for p_id, f_id in prod_forn_assoc:
            if p_id not in produto_fornecedores: produto_fornecedores[p_id] = []
            produto_fornecedores[p_id].append(fornecedores_map.get(f_id, ''))

        produto_naturezas = {}
        for p_id, n_id in prod_nat_assoc:
            if p_id not in produto_naturezas: produto_naturezas[p_id] = []
            produto_naturezas[p_id].append(naturezas_map.get(n_id, ''))

        produtos_json = []
        for produto in produtos_db:
            fornecedores_list = produto_fornecedores.get(produto.id_produto, [])
            naturezas_list = produto_naturezas.get(produto.id_produto, [])
            
            produtos_json.append({
                'id': produto.id_produto,
                'nome': produto.nome,
                'codigo': produto.codigo.strip() if produto.codigo else '',
                'descricao': produto.descricao,
                'preco': str(produto.preco),
                'codigoB': produto.codigoB,
                'codigoC': produto.codigoC,
                'fornecedores': ", ".join(sorted(fornecedores_list)),
                'naturezas': ", ".join(sorted(naturezas_list))
            })
            
        return jsonify(produtos_json), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos', methods=['POST'])
@jwt_required()
def add_novo_produto():
    try:
        dados = request.get_json()
        required_fields = ['nome', 'codigo']
        
        if not all(field in dados and dados[field] for field in required_fields):
            return jsonify({'erro': 'Campos obrigatórios (nome, codigo) não podem estar vazios.'}), 400

        novo_produto = Produto(
            nome=dados['nome'],
            codigo=dados['codigo'],
            descricao=dados.get('descricao'),
            preco=dados.get('preco', '0.00').replace(',', '.'), 
            codigoB=dados.get('codigoB'),
            codigoC=dados.get('codigoC')
        )
        db.session.add(novo_produto)
        db.session.commit()
        
        return jsonify({
            'mensagem': 'Produto adicionado com sucesso!',
            'id_produto_criado': novo_produto.id_produto
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/importar', methods=['POST'])
@jwt_required()
def importar_produtos_csv():
    if 'file' not in request.files:
        return jsonify({'erro': 'Nenhum ficheiro enviado.'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'erro': 'Nome de ficheiro vazio.'}), 400

    sucesso_count = 0
    erros = []
    
    try:
        file_bytes = file.stream.read()
        try:
            stream_content = file_bytes.decode("UTF-8")
        except UnicodeDecodeError:
            stream_content = file_bytes.decode("latin-1")
        
        stream = io.StringIO(stream_content, newline=None)
        header = stream.readline()
        stream.seek(0)
        delimiter = ';' if ';' in header else ','
        csv_reader = csv.DictReader(stream, delimiter=delimiter)

        id_usuario_logado = get_jwt_identity()

        for linha_num, linha in enumerate(csv_reader, start=2):
            try:
                codigo = linha.get('codigo', '').strip()
                nome = linha.get('nome', '').strip()
                preco_str = linha.get('preco', '0').strip()

                if not codigo or not nome:
                    erros.append(f"Linha {linha_num}: Campos obrigatórios (codigo, nome) em falta.")
                    continue

                if Produto.query.filter_by(codigo=codigo).first():
                    erros.append(f"Linha {linha_num}: Código '{codigo}' já existe.")
                    continue

                novo_produto = Produto(
                    codigo=codigo,
                    nome=nome,
                    preco=preco_str.replace(',', '.') if preco_str else '0.00',
                    descricao=linha.get('descricao', '').strip()
                )

                fornecedores_nomes = [fn.strip() for fn in linha.get('fornecedores_nomes', '').split(',') if fn.strip()]
                if fornecedores_nomes:
                    novo_produto.fornecedores.extend(Fornecedor.query.filter(Fornecedor.nome.in_(fornecedores_nomes)).all())

                naturezas_nomes = [nn.strip() for nn in linha.get('naturezas_nomes', '').split(',') if nn.strip()]
                if naturezas_nomes:
                    novo_produto.naturezas.extend(Natureza.query.filter(Natureza.nome.in_(naturezas_nomes)).all())

                db.session.add(novo_produto)
                db.session.flush()

                qtd_inicial = linha.get('quantidade', '0').strip()
                if qtd_inicial and int(qtd_inicial) > 0:
                    mov_inicial = MovimentacaoEstoque(
                        id_produto=novo_produto.id_produto,
                        id_usuario=id_usuario_logado,
                        quantidade=int(qtd_inicial),
                        tipo='Entrada',
                        motivo_saida='Balanço Inicial via Importação'
                    )
                    db.session.add(mov_inicial)
                
                sucesso_count += 1

            except Exception as e_interno:
                erros.append(f"Linha {linha_num}: Erro ao processar - {e_interno}.")
                continue

        db.session.commit()
        
        return jsonify({
            'mensagem': 'Importação concluída!',
            'produtos_importados': sucesso_count,
            'erros': erros
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': f'Erro geral ao processar: {str(e)}'}), 500

@app.route('/api/formularios/produto_data', methods=['GET'])
@jwt_required()
def get_form_produto_data():
    try:
        produto_id = request.args.get('produto_id', type=int)
        
        fornecedores_data = db.session.query(Fornecedor.id_fornecedor, Fornecedor.nome).order_by(Fornecedor.nome).all()
        naturezas_data = db.session.query(Natureza.id_natureza, Natureza.nome).order_by(Natureza.nome).all()
        
        dados_produto = None
        if produto_id:
            produto = Produto.query.options(
                joinedload(Produto.fornecedores),
                joinedload(Produto.naturezas)
            ).get(produto_id)
            
            if produto:
                dados_produto = {
                    'id': produto.id_produto,
                    'nome': produto.nome,
                    'codigo': produto.codigo.strip() if produto.codigo else '',
                    'descricao': produto.descricao,
                    'preco': str(produto.preco),
                    'codigoB': produto.codigoB,
                    'codigoC': produto.codigoC,
                    'fornecedores': [{'id': f.id_fornecedor} for f in produto.fornecedores],
                    'naturezas': [{'id': n.id_natureza} for n in produto.naturezas]
                }

        return jsonify({
            'fornecedores': [{'id': id, 'nome': nome} for id, nome in fornecedores_data],
            'naturezas': [{'id': id, 'nome': nome} for id, nome in naturezas_data],
            'produto': dados_produto
        }), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/<int:id_produto>', methods=['GET', 'PUT', 'DELETE'])
@jwt_required()
def produto_por_id_endpoint(id_produto):
    try:
        produto = Produto.query.get_or_404(id_produto)

        if request.method == 'GET':
            return jsonify({
                'id': produto.id_produto, 'nome': produto.nome,
                'codigo': produto.codigo.strip() if produto.codigo else '',
                'descricao': produto.descricao, 'preco': str(produto.preco),
                'codigoB': produto.codigoB, 'codigoC': produto.codigoC,
                'fornecedores': [{'id': f.id_fornecedor, 'nome': f.nome} for f in produto.fornecedores],
                'naturezas': [{'id': n.id_natureza, 'nome': n.nome} for n in produto.naturezas]
            }), 200
        
        elif request.method == 'PUT':
            dados = request.get_json()
            
            produto.nome = dados['nome']
            produto.codigo = dados['codigo']
            produto.descricao = dados.get('descricao')
            produto.preco = dados['preco']
            produto.codigoB = dados.get('codigoB')
            produto.codigoC = dados.get('codigoC')

            if 'fornecedores_ids' in dados:
                produto.fornecedores = []
                if dados['fornecedores_ids']:
                    produto.fornecedores = Fornecedor.query.filter(Fornecedor.id_fornecedor.in_(dados['fornecedores_ids'])).all()

            if 'naturezas_ids' in dados:
                produto.naturezas = []
                if dados['naturezas_ids']:
                    produto.naturezas = Natureza.query.filter(Natureza.id_natureza.in_(dados['naturezas_ids'])).all()

            db.session.commit()

            updated_product = Produto.query.options(joinedload(Produto.fornecedores), joinedload(Produto.naturezas)).get(id_produto)
            return jsonify({
                'id': updated_product.id_produto,
                'nome': updated_product.nome,
                'codigo': updated_product.codigo.strip(),
                'descricao': updated_product.descricao,
                'preco': str(updated_product.preco),
                'codigoB': updated_product.codigoB,
                'codigoC': updated_product.codigoC,
                'fornecedores': ", ".join(sorted([f.nome for f in updated_product.fornecedores])),
                'naturezas': ", ".join(sorted([n.nome for n in updated_product.naturezas]))
            }), 200
        
        elif request.method == 'DELETE':
            if MovimentacaoEstoque.query.filter_by(id_produto=id_produto).first():
                return jsonify({'erro': 'Produto possui histórico de movimentações e não pode ser excluído.'}), 400

            db.session.delete(produto)
            db.session.commit()
            return jsonify({'mensagem': 'Produto excluído com sucesso!'}), 200
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/codigo/<string:codigo>', methods=['GET'])
@jwt_required()
def get_produto_por_codigo(codigo):
    try:
        produto = Produto.query.filter_by(codigo=codigo.strip()).first()
        if produto:
            return jsonify({
                'id': produto.id_produto,
                'nome': produto.nome,
                'codigo': produto.codigo.strip(),
                'descricao': produto.descricao,
                'preco': str(produto.preco)
            }), 200
        return jsonify({'erro': 'Produto não encontrado.'}), 404
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/<int:id_produto>/estoque', methods=['GET'])
@jwt_required()
def get_saldo_estoque_produto(id_produto):
    try:
        Produto.query.get_or_404(id_produto)
        return jsonify({'id_produto': id_produto, 'saldo_atual': calcular_saldo_produto(id_produto)}), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500 

@app.route('/api/produtos/<int:id_produto>/fornecedores', methods=['POST'])
@jwt_required()
def adicionar_fornecedor_ao_produto(id_produto):
    try:
        dados = request.get_json()
        id_fornecedor = dados.get('id_fornecedor')
        if not id_fornecedor: return jsonify({'erro': 'ID do fornecedor obrigatório.'}), 400

        produto = Produto.query.get_or_404(id_produto)
        fornecedor = Fornecedor.query.get_or_404(id_fornecedor)

        if fornecedor not in produto.fornecedores:
            produto.fornecedores.append(fornecedor)
            db.session.commit()
        return jsonify({'mensagem': 'Associação realizada.'}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/<int:id_produto>/naturezas', methods=['POST'])
@jwt_required()
def adicionar_natureza_ao_produto(id_produto):
    try:
        dados = request.get_json()
        id_natureza = dados.get('id_natureza')
        if not id_natureza: return jsonify({'erro': 'ID da natureza obrigatório.'}), 400

        produto = Produto.query.get_or_404(id_produto)
        natureza = Natureza.query.get_or_404(id_natureza)

        if natureza not in produto.naturezas:
            produto.naturezas.append(natureza)
            db.session.commit()
        return jsonify({'mensagem': 'Associação realizada.'}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/<int:id_produto>/fornecedores/<int:id_fornecedor>', methods=['DELETE'])
@jwt_required()
def remover_fornecedor_do_produto(id_produto, id_fornecedor):
    try:
        produto = Produto.query.get_or_404(id_produto)
        fornecedor = Fornecedor.query.get_or_404(id_fornecedor)

        if fornecedor in produto.fornecedores:
            produto.fornecedores.remove(fornecedor)
            db.session.commit()
            return jsonify({'mensagem': 'Associação removida.'}), 200
        return jsonify({'erro': 'Associação não encontrada.'}), 404

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/produtos/<int:id_produto>/naturezas/<int:id_natureza>', methods=['DELETE'])
@jwt_required()
def remover_natureza_do_produto(id_produto, id_natureza):
    try:
        produto = Produto.query.get_or_404(id_produto)
        natureza = Natureza.query.get_or_404(id_natureza)

        if natureza in produto.naturezas:
            produto.naturezas.remove(natureza)
            db.session.commit()
            return jsonify({'mensagem': 'Associação removida.'}), 200
        return jsonify({'erro': 'Associação não encontrada.'}), 404

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

# ==============================================================================
# ROTAS: ESTOQUE
# ==============================================================================

@app.route('/api/estoque/entrada', methods=['POST'])
@jwt_required()
def registrar_entrada():
    try:
        dados = request.get_json()
        if 'id_produto' not in dados or 'quantidade' not in dados:
            return jsonify({'erro': 'Campos obrigatórios em falta'}), 400

        id_produto = dados['id_produto']
        qtd = dados['quantidade']
        
        saldo_atual = calcular_saldo_produto(id_produto)

        nova_entrada = MovimentacaoEstoque(
            id_produto=id_produto,
            quantidade=qtd,
            id_usuario=get_jwt_identity(),
            tipo='Entrada'
        )
        db.session.add(nova_entrada)
        db.session.commit()
        
        return jsonify({'mensagem': 'Entrada registrada!', 'novo_saldo': saldo_atual + qtd}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/estoque/saida', methods=['POST'])
@jwt_required()
def registrar_saida():
    try:
        dados = request.get_json()
        required = ['id_produto', 'quantidade', 'motivo_saida']
        if not all(k in dados for k in required):
            return jsonify({'erro': 'Campos obrigatórios em falta'}), 400

        id_produto = dados['id_produto']
        qtd = dados['quantidade']
        
        saldo_atual = calcular_saldo_produto(id_produto)
        if saldo_atual < qtd:
            return jsonify({'erro': f'Estoque insuficiente. Saldo atual: {saldo_atual}'}), 400

        nova_saida = MovimentacaoEstoque(
            id_produto=id_produto,
            quantidade=qtd,
            id_usuario=get_jwt_identity(),
            tipo='Saida',
            motivo_saida=dados.get('motivo_saida')
        )
        db.session.add(nova_saida)
        db.session.commit()
        
        return jsonify({'mensagem': 'Saída registrada!', 'novo_saldo': saldo_atual - qtd}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/estoque/saldos', methods=['GET'])
@jwt_required()
def get_saldos_estoque():
    try:
        termo = request.args.get('search')
        query = Produto.query

        if termo:
            query = query.filter(or_(
                Produto.nome.ilike(f"%{termo}%"),
                Produto.codigo.ilike(f"%{termo}%"),
                Produto.codigoB.ilike(f"%{termo}%"),
                Produto.codigoC.ilike(f"%{termo}%")
            ))

        produtos = query.all()
        saldos_json = []
        
        for p in produtos:
            saldos_json.append({
                'id_produto': p.id_produto,
                'codigo': p.codigo.strip() if p.codigo else '',
                'nome': p.nome,
                'saldo_atual': calcular_saldo_produto(p.id_produto),
                'preco': str(p.preco),
                'codigoB': p.codigoB.strip() if p.codigoB else '',
                'codigoC': p.codigoC.strip() if p.codigoC else ''
            })
            
        return jsonify(saldos_json), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/movimentacoes', methods=['GET'])
@jwt_required()
def get_todas_movimentacoes():
    try:
        filtro_tipo = request.args.get('tipo')
        query = MovimentacaoEstoque.query.options(
            joinedload(MovimentacaoEstoque.produto),
            joinedload(MovimentacaoEstoque.usuario)
        ).order_by(MovimentacaoEstoque.data_hora.desc())

        if filtro_tipo in ["Entrada", "Saida"]:
            query = query.filter(MovimentacaoEstoque.tipo == filtro_tipo)

        movimentacoes = query.all()
        resultado = []
        for mov in movimentacoes:
            resultado.append({
                'id': mov.id_movimentacao,
                'data_hora': mov.data_hora.strftime('%d/%m/%Y %H:%M:%S'),
                'tipo': mov.tipo,
                'quantidade': mov.quantidade,
                'motivo_saida': mov.motivo_saida,
                'produto_codigo': mov.produto.codigo.strip() if mov.produto else 'N/A',
                'produto_nome': mov.produto.nome if mov.produto else 'Produto Excluído',
                'usuario_nome': mov.usuario.nome if mov.usuario else 'Usuário Excluído'
            })
        return jsonify(resultado), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

# ==============================================================================
# ROTAS: AUTENTICAÇÃO E USUÁRIOS
# ==============================================================================

@app.route('/api/login', methods=['POST'])
def login_endpoint():
    try:
        dados = request.get_json()
        if not dados or 'login' not in dados or 'senha' not in dados:
            return jsonify({'erro': 'Login e senha obrigatórios'}), 400

        usuario = Usuario.query.filter_by(login=dados['login'], ativo=True).first()

        if usuario and usuario.check_password(dados['senha']):
            access_token = create_access_token(
                identity=str(usuario.id_usuario),
                additional_claims={'permissao': usuario.permissao},
                expires_delta=timedelta(hours=8)
            )
            return jsonify(access_token=access_token), 200
        else:
            return jsonify({"erro": "Credenciais inválidas"}), 401
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/usuario/me', methods=['GET'])
@jwt_required()
def get_usuario_logado():
    usuario = Usuario.query.get(get_jwt_identity())
    if not usuario: return jsonify({"erro": "Usuário não encontrado"}), 404
    
    return jsonify({
        'id': usuario.id_usuario,
        'nome': usuario.nome,
        'login': usuario.login,
        'permissao': usuario.permissao
    }), 200

@app.route('/api/usuario/mudar-senha', methods=['POST'])
@jwt_required()
def mudar_senha_usuario():
    try:
        usuario = Usuario.query.get(get_jwt_identity())
        dados = request.get_json()
        
        senha_atual = dados.get('senha_atual')
        nova_senha = dados.get('nova_senha')
        confirmacao = dados.get('confirmacao_nova_senha')

        if not all([senha_atual, nova_senha, confirmacao]):
            return jsonify({'erro': 'Todos os campos são obrigatórios.'}), 400

        if not usuario.check_password(senha_atual):
            return jsonify({'erro': 'Senha atual incorreta.'}), 401

        if nova_senha != confirmacao:
            return jsonify({'erro': 'Nova senha e confirmação divergem.'}), 400
            
        usuario.set_password(nova_senha)
        db.session.commit()
        return jsonify({'mensagem': 'Senha alterada com sucesso!'}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500 

@app.route('/api/usuarios', methods=['GET'])
@jwt_required()
def get_todos_usuarios():
    try:
        if get_jwt().get('permissao') != 'Administrador':
            return jsonify({"erro": "Acesso negado."}), 403

        usuarios = Usuario.query.all()
        return jsonify([{
            'id': u.id_usuario,
            'nome': u.nome,
            'login': u.login,
            'permissao': u.permissao,
            'ativo': u.ativo
        } for u in usuarios]), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/usuarios', methods=['POST'])
@jwt_required()
def add_novo_usuario():
    try:
        if get_jwt().get('permissao') != 'Administrador':
            return jsonify({"erro": "Acesso negado."}), 403

        dados = request.get_json()
        if not all(k in dados for k in ['login', 'senha', 'nome', 'permissao']):
            return jsonify({'erro': 'Dados incompletos'}), 400

        novo_usuario = Usuario(
            nome=dados['nome'],
            login=dados['login'],
            permissao=dados['permissao']
        )
        novo_usuario.set_password(dados['senha'])
        
        db.session.add(novo_usuario)
        db.session.commit()
        return jsonify({'mensagem': 'Usuário criado!'}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/usuarios/<int:id_usuario>', methods=['GET', 'PUT', 'DELETE'])
@jwt_required()
def usuario_por_id_endpoint(id_usuario):
    if get_jwt().get('permissao') != 'Administrador':
        return jsonify({"erro": "Acesso negado."}), 403
    
    try:
        usuario = Usuario.query.get_or_404(id_usuario)

        if request.method == 'GET':
            return jsonify({
                'id': usuario.id_usuario, 'nome': usuario.nome,
                'login': usuario.login, 'permissao': usuario.permissao, 'ativo': usuario.ativo
            }), 200

        elif request.method == 'PUT':
            dados = request.get_json()
            usuario.nome = dados.get('nome', usuario.nome)
            usuario.login = dados.get('login', usuario.login)
            usuario.permissao = dados.get('permissao', usuario.permissao)
            
            if dados.get('senha'):
                usuario.set_password(dados['senha'])
            
            db.session.commit()
            return jsonify({'mensagem': 'Usuário atualizado!'}), 200

        elif request.method == 'DELETE':
            usuario.ativo = not usuario.ativo
            db.session.commit()
            status = "desativado" if not usuario.ativo else "reativado"
            return jsonify({'mensagem': f'Usuário {status}!'}), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

# ==============================================================================
# ROTAS: TABELAS AUXILIARES (FORNECEDOR/NATUREZA)
# ==============================================================================

@app.route('/api/fornecedores', methods=['GET', 'POST'])
@jwt_required()
def gerir_fornecedores():
    if request.method == 'GET':
        items = Fornecedor.query.order_by(Fornecedor.nome).all()
        return jsonify([{'id': i.id_fornecedor, 'nome': i.nome} for i in items]), 200
    
    try:
        dados = request.get_json()
        if not dados.get('nome'): return jsonify({'erro': 'Nome obrigatório'}), 400
        
        db.session.add(Fornecedor(nome=dados['nome']))
        db.session.commit()
        return jsonify({'mensagem': 'Fornecedor criado!'}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/fornecedores/<int:id_fornecedor>', methods=['GET', 'PUT', 'DELETE'])
@jwt_required()
def fornecedor_id_ops(id_fornecedor):
    try:
        fornecedor = Fornecedor.query.get_or_404(id_fornecedor)
        
        if request.method == 'GET':
            return jsonify({'id': fornecedor.id_fornecedor, 'nome': fornecedor.nome}), 200
            
        elif request.method == 'PUT':
            dados = request.get_json()
            if not dados.get('nome'): return jsonify({'erro': 'Nome obrigatório'}), 400
            fornecedor.nome = dados['nome']
            db.session.commit()
            return jsonify({'mensagem': 'Atualizado!'}), 200
            
        elif request.method == 'DELETE':
            if fornecedor.produtos:
                return jsonify({'erro': 'Possui associações. Não pode excluir.'}), 400
            db.session.delete(fornecedor)
            db.session.commit()
            return jsonify({'mensagem': 'Excluído!'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/naturezas', methods=['GET', 'POST'])
@jwt_required()
def gerir_naturezas():
    if request.method == 'GET':
        items = Natureza.query.order_by(Natureza.nome).all()
        return jsonify([{'id': i.id_natureza, 'nome': i.nome} for i in items]), 200
    
    try:
        dados = request.get_json()
        if not dados.get('nome'): return jsonify({'erro': 'Nome obrigatório'}), 400
        
        db.session.add(Natureza(nome=dados['nome']))
        db.session.commit()
        return jsonify({'mensagem': 'Natureza criada!'}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/naturezas/<int:id_natureza>', methods=['GET', 'PUT', 'DELETE'])
@jwt_required()
def natureza_id_ops(id_natureza):
    try:
        natureza = Natureza.query.get_or_404(id_natureza)
        
        if request.method == 'GET':
            return jsonify({'id': natureza.id_natureza, 'nome': natureza.nome}), 200
            
        elif request.method == 'PUT':
            dados = request.get_json()
            if not dados.get('nome'): return jsonify({'erro': 'Nome obrigatório'}), 400
            natureza.nome = dados['nome']
            db.session.commit()
            return jsonify({'mensagem': 'Atualizado!'}), 200
            
        elif request.method == 'DELETE':
            if natureza.produtos:
                return jsonify({'erro': 'Possui associações. Não pode excluir.'}), 400
            db.session.delete(natureza)
            db.session.commit()
            return jsonify({'mensagem': 'Excluído!'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

# ==============================================================================
# ROTAS: DOCUMENTOS E DASHBOARD
# ==============================================================================

@app.route('/api/dashboard/kpis', methods=['GET'])
@jwt_required()
def get_dashboard_kpis():
    try:
        total_produtos = db.session.query(func.count(Produto.id_produto)).scalar()
        total_fornecedores = db.session.query(func.count(Fornecedor.id_fornecedor)).scalar()

        subquery_saldos = db.session.query(
            MovimentacaoEstoque.id_produto,
            func.sum(case(
                (MovimentacaoEstoque.tipo == 'Entrada', MovimentacaoEstoque.quantidade),
                (MovimentacaoEstoque.tipo == 'Saida', -MovimentacaoEstoque.quantidade)
            )).label('saldo')
        ).group_by(MovimentacaoEstoque.id_produto).subquery()

        valor_total_estoque = db.session.query(
            func.sum(Produto.preco * subquery_saldos.c.saldo)
        ).join(subquery_saldos, Produto.id_produto == subquery_saldos.c.id_produto).scalar() or 0

        return jsonify({
            'total_produtos': total_produtos,
            'total_fornecedores': total_fornecedores,
            'valor_total_estoque': float(valor_total_estoque)
        }), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/servicos/<int:servico_id>/documentos', methods=['GET'])
@jwt_required()
def get_historico_documentos(servico_id):
    try:
        documentos = DocumentosGerados.query.options(joinedload(DocumentosGerados.usuario))\
            .filter_by(servico_id=servico_id)\
            .order_by(DocumentosGerados.versao.desc()).all()

        historico_list = [{
            'id': doc.id,
            'data_criacao': doc.data_criacao.strftime('%d/%m/%Y'),
            'versao': doc.versao,
            'caminho_pdf': doc.caminho_pdf_final,
            'nome_usuario': doc.usuario.nome if doc.usuario else 'Desconhecido'
        } for doc in documentos]
        
        return jsonify(historico_list), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

@app.route('/api/documentos/<int:documento_id>', methods=['GET', 'DELETE'])
@jwt_required()
def gerir_documento_por_id(documento_id):
    try:
        if request.method == 'DELETE':
            if get_jwt().get('permissao') != 'Administrador':
                return jsonify({"erro": "Acesso negado."}), 403

            documento = DocumentosGerados.query.get_or_404(documento_id)
            if os.path.exists(documento.caminho_pdf_final):
                os.remove(documento.caminho_pdf_final)

            db.session.delete(documento)
            db.session.commit()
            return jsonify({'mensagem': 'Documento excluído.'}), 200

        documento = DocumentosGerados.query.get_or_404(documento_id)
        return jsonify(documento.dados_formulario), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'erro': str(e)}), 500

@app.route('/api/servicos/<int:servico_id>/documentos', methods=['POST'])
@jwt_required()
def gerar_novo_documento(servico_id):
    dir_path = os.path.dirname(os.path.realpath(__file__))
    temp_docx = os.path.join(dir_path, 'temp_documento.docx')
    temp_pdf = os.path.join(dir_path, 'temp_documento.pdf')
    output_folder = os.path.join(dir_path, 'documentos_gerados')

    try:
        if 'dados_formulario' not in request.form:
            return jsonify({'erro': 'Dados ausentes'}), 400
        
        dados_formulario_str = request.form.get('dados_formulario')
        dados_formulario = json.loads(dados_formulario_str)
        anexos = request.files.getlist('anexos')
        id_usuario_logado = get_jwt_identity()
        
        versao_anterior = db.session.query(func.max(DocumentosGerados.versao)).filter_by(servico_id=servico_id).scalar()
        nova_versao = (versao_anterior or 0) + 1

        doc = Document(os.path.join(dir_path, 'template.docx'))
        
        replacements = {}
        replacements.update(dados_formulario.get('identificacao_projeto', {}))
        replacements.update(dados_formulario.get('escopo_premissas', {}))
        replacements.update(dados_formulario.get('diagramas_desenhos', {}))
        replacements.update(dados_formulario.get('testes_comissionamento', {}))
        replacements.update(dados_formulario.get('operacao_manutencao', {}))
        replacements.update(dados_formulario.get('treinamento', {}))
        replacements.update(dados_formulario.get('documentos_as_built', {}))
        replacements.update(dados_formulario.get('anexos', {}))
        
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if f'{{{{{key}}}}}' in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(f'{{{{{key}}}}}', str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if f'{{{{{key}}}}}' in p.text:
                                for run in p.runs:
                                    run.text = run.text.replace(f'{{{{{key}}}}}', str(value))

        try:
            if 'lista_documentos_projeto' in dados_formulario:
                tabela_docs = doc.tables[0]
                for item in dados_formulario['lista_documentos_projeto']:
                    celulas = tabela_docs.add_row().cells
                    celulas[0].text, celulas[1].text, celulas[2].text, celulas[3].text, celulas[4].text, celulas[5].text = item.get('titulo', ''), item.get('codigo', ''), item.get('revisao', ''), item.get('data', ''), item.get('autor', ''), item.get('status', '')

            if 'lista_instrumentos' in dados_formulario:
                tabela_instrumentos = doc.tables[1]
                for item in dados_formulario['lista_instrumentos']:
                    celulas = tabela_instrumentos.add_row().cells
                    celulas[0].text, celulas[1].text, celulas[2].text, celulas[3].text, celulas[4].text, celulas[5].text = item.get('tag', ''), item.get('descricao', ''), item.get('fabricante_modelo', ''), item.get('faixa', ''), item.get('sinal', ''), item.get('localizacao', '')

            if 'programacao_logica' in dados_formulario:
                tabela_programacao = doc.tables[2]
                for item in dados_formulario['programacao_logica']:
                    celulas = tabela_programacao.add_row().cells
                    celulas[0].text, celulas[1].text = item.get('ficheiro', ''), item.get('descricao', '')

            if 'treinamento' in dados_formulario and 'participantes' in dados_formulario['treinamento']:
                tabela_participantes = doc.tables[3]
                for item in dados_formulario['treinamento']['participantes']:
                    celulas = tabela_participantes.add_row().cells
                    celulas[0].text, celulas[1].text = item.get('nome', ''), item.get('certificado', '')

            if 'documentos_as_built' in dados_formulario:
                tabela_as_built = doc.tables[4]
                for item in dados_formulario['documentos_as_built']:
                    celulas = tabela_as_built.add_row().cells
                    celulas[0].text, celulas[1].text = item.get('documento', ''), item.get('notas', '')
        except IndexError:
            print("AVISO: Tabelas insuficientes no template.")

        doc.save(temp_docx)

        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", dir_path, temp_docx], check=True)
        
        merger = PdfWriter()
        merger.append(temp_pdf)
        for anexo in anexos:
            merger.append(anexo.stream)
        
        filename = f"servico_{servico_id}_v{nova_versao}.pdf"
        filepath = os.path.join(output_folder, filename)
        with open(filepath, "wb") as f_out:
            merger.write(f_out)

        db.session.add(DocumentosGerados(
            servico_id=servico_id,
            usuario_id=id_usuario_logado,
            versao=nova_versao,
            dados_formulario=dados_formulario,
            caminho_pdf_final=filepath
        ))
        db.session.commit()

        return jsonify({'mensagem': 'Gerado com sucesso', 'caminho_download': filepath}), 201

    except Exception as e:
        db.session.rollback()
        traceback.print_exc()
        return jsonify({'erro': str(e)}), 500
    finally:
        if os.path.exists(temp_docx): os.remove(temp_docx)
        if os.path.exists(temp_pdf): os.remove(temp_pdf)

@app.route('/api/versao', methods=['GET'])
def get_versao_app():
    try:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'versao.json')
        with open(path, 'r', encoding="utf-8") as f:
            return jsonify(json.load(f)), 200
    except Exception as e:
        return jsonify({'erro': str(e)}), 500

# ==============================================================================
# GERADORES DE PDF
# ==============================================================================

def gerar_pdf_etiquetas(produtos):
    buffer = io.BytesIO()
    w, h = 62 * mm, 100 * mm
    doc = SimpleDocTemplate(buffer, pagesize=(w, h), leftMargin=5*mm, rightMargin=5*mm, topMargin=5*mm, bottomMargin=5*mm)
    
    styles = getSampleStyleSheet()
    styles['Normal'].fontSize = 12
    styles['Normal'].leading = 14
    style_cod = ParagraphStyle(name='CodigoStyle', parent=styles['Normal'], alignment=TA_CENTER)
    
    elems = []
    for p in produtos:
        elems.append(Paragraph(f"<b>{p.nome}</b>", styles['Normal']))
        elems.append(Spacer(1, 8*mm))
        elems.append(code128.Code128(p.codigo, barHeight=20*mm, barWidth=0.4*mm))
        elems.append(Spacer(1, 2*mm))
        elems.append(Paragraph(p.codigo, style_cod))
        elems.append(PageBreak())
        
    if elems: elems.pop()
    doc.build(elems)
    buffer.seek(0)
    return buffer

def gerar_inventario_pdf(dados):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    elems = []
    styles = getSampleStyleSheet()

    elems.append(Paragraph("Relatório de Inventário Atual", styles['h1']))
    elems.append(Spacer(1, 12))

    table_data = [["Código", "Nome", "Saldo", "Preço Unit. (R$)", "Valor Total (R$)"]]
    total_geral = 0
    
    for item in dados:
        total_item = item['saldo_atual'] * item['preco']
        total_geral += total_item
        table_data.append([
            item['codigo'], item['nome'], str(item['saldo_atual']),
            f"{float(item['preco']):.2f}", f"{float(total_item):.2f}"
        ])

    t = Table(table_data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ]))
    elems.append(t)
    elems.append(Spacer(1, 12))
    elems.append(Paragraph(f"<b>Valor Total do Estoque:</b> R$ {float(total_geral):.2f}", styles['h3']))
    
    doc.build(elems)
    buffer.seek(0)
    return buffer

def gerar_historico_pdf(dados):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elems = []
    styles = getSampleStyleSheet()

    elems.append(Paragraph("Relatório de Histórico de Movimentações", styles['h1']))
    elems.append(Spacer(1, 12))
    
    table_data = [["Data/Hora", "Produto", "Tipo", "Qtd.", "Saldo Após", "Usuário", "Motivo"]]
    for i in dados:
        table_data.append([
            i['data_hora'], f"{i['produto_codigo']} - {i['produto_nome']}",
            i['tipo'], str(i['quantidade']), str(i['saldo_apos']),
            i['usuario_nome'], i.get('motivo_saida', '')
        ])

    t = Table(table_data, colWidths=[110, 180, 50, 40, 60, 100, 130])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ]))
    elems.append(t)
    doc.build(elems)
    buffer.seek(0)
    return buffer

# ==============================================================================
# ENDPOINTS RELATÓRIOS
# ==============================================================================

@app.route('/api/relatorios/inventario', methods=['GET'])
@jwt_required()
def relatorio_inventario():
    formato = request.args.get('formato', 'pdf').lower()
    produtos = Produto.query.all()
    dados_relatorio = []
    for produto in produtos:
        saldo = calcular_saldo_produto(produto.id_produto)
        dados_relatorio.append({
            'codigo': produto.codigo.strip(),
            'nome': produto.nome,
            'saldo_atual': saldo,
            'preco': produto.preco
        })

    if formato == 'xlsx':
        df = pd.DataFrame(dados_relatorio)
        df['valor_total'] = df['saldo_atual'] * df['preco']
        df = df.rename(columns={'codigo': 'Código', 'nome': 'Nome', 'saldo_atual': 'Saldo', 'preco': 'Preço Unitário (R$)', 'valor_total': 'Valor Total (R$)'})
        
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False, engine='openpyxl')
        buffer.seek(0)
        return send_file(buffer, download_name="relatorio_inventario.xlsx", as_attachment=True)
    
    pdf_buffer = gerar_inventario_pdf(dados_relatorio)
    return send_file(pdf_buffer, download_name="relatorio_inventario.pdf", as_attachment=True)

@app.route('/api/relatorios/movimentacoes', methods=['GET'])
@jwt_required()
def relatorio_movimentacoes():
    formato = request.args.get('formato', 'json').lower()
    data_inicio_str = request.args.get('data_inicio')
    data_fim_str = request.args.get('data_fim')
    tipo = request.args.get('tipo')

    query = MovimentacaoEstoque.query.options(
        joinedload(MovimentacaoEstoque.produto),
        joinedload(MovimentacaoEstoque.usuario)
    ).order_by(MovimentacaoEstoque.id_produto, MovimentacaoEstoque.data_hora)

    if data_inicio_str:
        query = query.filter(MovimentacaoEstoque.data_hora >= datetime.strptime(data_inicio_str, '%Y-%m-%d'))
    if data_fim_str:
        data_fim = datetime.strptime(data_fim_str, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
        query = query.filter(MovimentacaoEstoque.data_hora <= data_fim)
    
    todas_movimentacoes = query.all()

    dados_relatorio = []
    saldos_atuais = {}

    for mov in todas_movimentacoes:
        id_produto = mov.id_produto
        
        if id_produto not in saldos_atuais:
            saldo_inicial_query = db.session.query(
                func.sum(case(
                    (MovimentacaoEstoque.tipo == 'Entrada', MovimentacaoEstoque.quantidade),
                    (MovimentacaoEstoque.tipo == 'Saida', -MovimentacaoEstoque.quantidade)
                ))
            ).filter(
                MovimentacaoEstoque.id_produto == id_produto,
                MovimentacaoEstoque.data_hora < datetime.strptime(data_inicio_str, '%Y-%m-%d') if data_inicio_str else True
            )
            saldo_inicial = saldo_inicial_query.scalar() or 0
            saldos_atuais[id_produto] = saldo_inicial

        if mov.tipo == 'Entrada':
            saldos_atuais[id_produto] += mov.quantidade
        else:
            saldos_atuais[id_produto] -= mov.quantidade
        
        dados_relatorio.append({
            'data_hora': mov.data_hora.strftime('%d/%m/%Y %H:%M:%S'),
            'produto_codigo': mov.produto.codigo.strip() if mov.produto else 'N/A',
            'produto_nome': mov.produto.nome if mov.produto else 'Produto Excluído',
            'tipo': mov.tipo,
            'quantidade': mov.quantidade,
            'saldo_apos': saldos_atuais[id_produto],
            'usuario_nome': mov.usuario.nome if mov.usuario else 'Usuário Excluído',
            'motivo_saida': mov.motivo_saida if mov.motivo_saida else ''
        })

    if tipo and tipo in ["Entrada", "Saida"]:
        dados_relatorio = [linha for linha in dados_relatorio if linha['tipo'] == tipo]

    dados_relatorio.sort(key=lambda x: datetime.strptime(x['data_hora'], '%d/%m/%Y %H:%M:%S'), reverse=True)

    if formato == 'json':
        return jsonify(dados_relatorio), 200
    
    elif formato == 'xlsx':
        df = pd.DataFrame(dados_relatorio)
        df = df.rename(columns={
            'data_hora': 'Data/Hora', 'produto_codigo': 'Cód. Produto', 'produto_nome': 'Nome Produto',
            'tipo': 'Tipo', 'quantidade': 'Qtd. Mov.', 'saldo_apos': 'Saldo Após', 'usuario_nome': 'Usuário', 'motivo_saida': 'Motivo da Saída'
        })
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False, engine='openpyxl')
        buffer.seek(0)
        return send_file(buffer, download_name="relatorio_movimentacoes.xlsx", as_attachment=True)
        
    pdf_buffer = gerar_historico_pdf(dados_relatorio)
    return send_file(pdf_buffer, download_name="relatorio_movimentacoes.pdf", as_attachment=True)

@app.route('/api/produtos/etiquetas', methods=['POST'])
@jwt_required()
def gerar_etiquetas_produtos():
    try:
        dados = request.get_json()
        if not dados or 'product_ids' not in dados:
            return jsonify({'erro': 'Lista de IDs de produtos em falta.'}), 400

        produtos = Produto.query.filter(Produto.id_produto.in_(dados['product_ids'])).all()
        if not produtos:
            return jsonify({'erro': 'Nenhum produto encontrado.'}), 404

        pdf_buffer = gerar_pdf_etiquetas(produtos)
        return send_file(pdf_buffer, as_attachment=True, download_name="etiquetas.pdf", mimetype='application/pdf')

    except Exception as e:
        return jsonify({'erro': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)